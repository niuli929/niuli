"""
批量抖音博主文案提取流水线
用法: python batch_pipeline.py --creator "博主名" --sec-uid "SEC_UID" --batch-size 10 [--start-batch 1]
"""

import os, sys, json, time, subprocess, re, argparse
from pathlib import Path
from datetime import datetime

# ===================== 配置 =====================
WORK_DIR = Path(r'C:\工作区\Claw')
DY_CLI = r'C:\Users\niu\.workbuddy\binaries\python\versions\3.13.12\Scripts\dy.exe'
FFMPEG = r'C:\Users\niu\Tools\ffmpeg\ffmpeg-8.1.1-essentials_build\bin\ffmpeg.exe'
FFPROBE = r'C:\Users\niu\Tools\ffmpeg\ffmpeg-8.1.1-essentials_build\bin\ffprobe.exe'
PROGRESS_FILE = WORK_DIR / 'batch_progress.json'
MIN_AUDIO_SECONDS = 5  # 低于此值视为图文视频，跳过

# IMA 配置
KB_ID = 'ZDDkqbhVJ_cBBzicExDbTgZTvZfKa6VCXkPTP57OEhE='
IMA_API_SCRIPT = r'C:\Users\niu\.workbuddy\skills\skill_2053082144792322048\ima_api.cjs'
IMA_COS_UPLOAD = r'C:\Users\niu\.workbuddy\skills\skill_2053082144792322048\knowledge-base\scripts\cos-upload.cjs'
IMA_OPTS = "{\\"clientId\\": \\"your-ima-client-id\\", \\"apiKey\\": \\"your-ima-api-key\\"}"


def run(cmd, timeout=120):
    """运行命令，返回 (stdout, stderr, returncode)"""
    try:
        r = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=timeout)
        return r.stdout.strip(), r.stderr.strip(), r.returncode
    except subprocess.TimeoutExpired:
        return '', f'Timeout after {timeout}s', -1


def load_progress():
    if PROGRESS_FILE.exists():
        with open(PROGRESS_FILE) as f:
            return json.load(f)
    return {}


def save_progress(prog):
    with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
        json.dump(prog, f, ensure_ascii=False, indent=2)


def download_batch(sec_uid, limit, output_dir, batch_num):
    """下载一批视频"""
    out, err, rc = run(
        f'"{DY_CLI}" download --user "{sec_uid}" --limit {limit} -o "{output_dir}"',
        timeout=180
    )
    print(f'  dy-cli 返回码: {rc}')
    if '未找到作品' in out:
        print('  ⚠ 未找到作品')
        return []
    
    # 找到下载目录（dy-cli 会创建 博主名 子目录）
    creator_dir = None
    for d in Path(output_dir).iterdir():
        if d.is_dir() and not d.name.startswith('.'):
            creator_dir = d
            break
    
    if not creator_dir:
        print('  ⚠ 未找到下载目录')
        return []
    
    videos = sorted(creator_dir.glob('*.mp4'))
    print(f'  下载目录: {creator_dir}')
    print(f'  已下载 {len(videos)} 个视频')
    return videos


def get_audio_duration(wav_path):
    """获取音频时长（秒）"""
    out, _, rc = run(
        f'"{FFPROBE}" -v error -show_entries format=duration -of default=noprint_wrappers=1:nokey=1 "{wav_path}"',
        timeout=15
    )
    if rc == 0 and out:
        try:
            return float(out.strip())
        except:
            pass
    return 0


def extract_audio(video_path, wav_path):
    """提取音频并检查时长"""
    out, err, rc = run(
        f'"{FFMPEG}" -y -i "{video_path}" -vn -ac 1 -ar 16000 -sample_fmt s16 "{wav_path}"',
        timeout=120
    )
    if rc != 0:
        print(f'    ffmpeg 失败: {err[:200]}')
        return 0
    
    duration = get_audio_duration(wav_path)
    return duration


def transcribe_audio(wav_path, txt_path, model):
    """转写音频"""
    segments, info = model.transcribe(
        str(wav_path), language='zh', beam_size=5,
        vad_filter=True, vad_parameters={'min_silence_duration_ms': 500}
    )
    lines = []
    with open(txt_path, 'w', encoding='utf-8') as f:
        for seg in segments:
            text = seg.text.strip()
            if text and len(text) >= 2:
                lines.append(text)
                f.write(text + '\n')
    return lines, info.language, info.language_probability


def generate_word(txt_path, docx_path, title, creator):
    """根据转写文本生成精简版 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 读取原文
    with open(txt_path, encoding='utf-8') as f:
        raw_text = f.read()
    
    # 基础清理：去语气词、拼句
    lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
    cleaned = '\n'.join(lines)
    
    doc = Document()
    
    # 标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 来源
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'来源：抖音@{creator}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_paragraph()
    
    # 正文：按段落分块
    paragraphs = [p for p in cleaned.split('\n') if len(p) > 5]
    for para in paragraphs[:50]:  # 最多50段
        doc.add_paragraph(para)
    
    doc.save(str(docx_path))
    return len(paragraphs)


def upload_to_ima(file_path, file_name):
    """上传文件到 IMA 知识库（使用 subprocess list 避免 shell 转义问题）"""
    
    file_size = os.path.getsize(file_path)
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == '.docx':
        media_type = 3
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        media_type = 13
        content_type = 'text/plain'
    
    api_dir = os.path.dirname(IMA_API_SCRIPT)
    cos_dir = os.path.dirname(IMA_COS_UPLOAD)
    
    # Step 1: check duplicates
    dup_payload = json.dumps({
        'params': [{'name': file_name, 'media_type': media_type}],
        'knowledge_base_id': KB_ID
    }, ensure_ascii=False)
    r = subprocess.run(
        ['node', IMA_API_SCRIPT, 'openapi/wiki/v1/check_repeated_names', dup_payload, IMA_OPTS],
        capture_output=True, text=True, timeout=15, cwd=api_dir
    )
    try:
        dup_resp = json.loads(r.stdout)
        if dup_resp.get('code') != 0:
            return False, f'check_repeated_names: {r.stdout[:200]}'
    except:
        return False, f'check_repeated_names JSON parse: {r.stdout[:200]}'
    
    # Step 2: create media
    cm_payload = json.dumps({
        'file_name': file_name, 'file_size': file_size,
        'content_type': content_type, 'knowledge_base_id': KB_ID,
        'file_ext': ext[1:]
    }, ensure_ascii=False)
    r = subprocess.run(
        ['node', IMA_API_SCRIPT, 'openapi/wiki/v1/create_media', cm_payload, IMA_OPTS],
        capture_output=True, text=True, timeout=15, cwd=api_dir
    )
    try:
        cm_resp = json.loads(r.stdout)
        if cm_resp.get('code') != 0:
            return False, f'create_media: {r.stdout[:200]}'
    except:
        return False, f'create_media JSON parse: {r.stdout[:200]}'
    
    media_id = cm_resp['data']['media_id']
    cred = cm_resp['data']['cos_credential']
    
    # Step 3: COS upload
    cos_args = [
        'node', IMA_COS_UPLOAD,
        '--file', str(file_path),
        '--secret-id', cred['secret_id'],
        '--secret-key', cred['secret_key'],
        '--token', cred['token'],
        '--bucket', cred['bucket_name'],
        '--region', cred['region'],
        '--cos-key', cred['cos_key'],
        '--content-type', content_type,
        '--start-time', str(cred['start_time']),
        '--expired-time', str(cred['expired_time']),
        '--timeout', '60000'
    ]
    r = subprocess.run(cos_args, capture_output=True, text=True, timeout=90, cwd=cos_dir)
    if 'Upload successful' not in r.stdout:
        return False, f'COS upload: {r.stdout[:200]}'
    
    # Step 4: add knowledge
    ak_payload = json.dumps({
        'media_type': media_type, 'media_id': media_id,
        'title': file_name, 'knowledge_base_id': KB_ID,
        'file_info': {'cos_key': cred['cos_key'], 'file_size': file_size, 'file_name': file_name}
    }, ensure_ascii=False)
    r = subprocess.run(
        ['node', IMA_API_SCRIPT, 'openapi/wiki/v1/add_knowledge', ak_payload, IMA_OPTS],
        capture_output=True, text=True, timeout=15, cwd=api_dir
    )
    try:
        ak_resp = json.loads(r.stdout)
        if ak_resp.get('code') != 0:
            return False, f'add_knowledge: {r.stdout[:200]}'
    except:
        return False, f'add_knowledge JSON parse: {r.stdout[:200]}'
    
    return True, media_id


def process_batch(sec_uid, creator_name, batch_size, batch_num, output_base):
    """处理一批视频"""
    import faster_whisper
    
    print(f'\n{"="*60}')
    print(f'  批次 {batch_num}: 下载 {batch_size} 条')
    print(f'  博主: {creator_name}')
    print(f'{"="*60}')
    
    output_dir = output_base / f'batch_{batch_num:03d}'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Step 1: 下载
    print('\n[1/5] 下载视频...')
    t0 = time.time()
    videos = download_batch(sec_uid, batch_size, str(output_dir), batch_num)
    if not videos:
        print('  没有下载到视频，终止')
        return {'downloaded': 0, 'skipped': 0, 'transcribed': 0, 'uploaded': 0, 'total_time': 0}
    print(f'  耗时 {time.time()-t0:.0f}s')
    
    # Step 2: 提取音频 + 检查时长
    print(f'\n[2/5] 提取音频 ({len(videos)} 个)...')
    audio_tasks = []
    for v in videos:
        wav = v.with_suffix('.wav')
        dur = extract_audio(v, wav)
        audio_tasks.append({
            'video': v,
            'wav': wav,
            'duration': dur,
            'skip': dur < MIN_AUDIO_SECONDS,
            'skip_reason': f'音频仅 {dur:.0f}s (阈值 {MIN_AUDIO_SECONDS}s)' if dur < MIN_AUDIO_SECONDS else None
        })
    
    skipped = [t for t in audio_tasks if t['skip']]
    to_transcribe = [t for t in audio_tasks if not t['skip']]
    
    print(f'  跳过 {len(skipped)} 条（图文/短音频）')
    for s in skipped:
        print(f'    ⊘ {s["video"].name}: {s["skip_reason"]}')
    print(f'  待转写 {len(to_transcribe)} 条')
    
    # Step 3: 转写
    print(f'\n[3/5] Whisper 转写 ({len(to_transcribe)} 条)...')
    os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
    model = faster_whisper.WhisperModel('small', device='cpu', compute_type='int8')
    
    transcribed = []
    for task in to_transcribe:
        v = task['video']
        wav = task['wav']
        label = v.stem[:6]
        txt_path = v.with_suffix('.txt')
        
        print(f'  转写: {v.name[:60]}... ({task["duration"]:.0f}s)')
        t1 = time.time()
        lines, lang, prob = transcribe_audio(wav, txt_path, model)
        elapsed = time.time() - t1
        print(f'    完成: {len(lines)} 句, {elapsed:.0f}s, 语言={lang}({prob:.2f})')
        transcribed.append({**task, 'lines': lines, 'txt_path': txt_path})
    
    # Step 4: 生成 Word
    print(f'\n[4/5] 生成 Word 文档 ({len(transcribed)} 条)...')
    docs = []
    for task in transcribed:
        txt = task['txt_path']
        # 用视频标题（去掉编号前缀）作为文档标题
        vname = task['video'].stem
        # 去掉 NNN_ 前缀
        title = re.sub(r'^\d{3,4}_', '', vname)
        # 截断过长标题
        if len(title) > 60:
            title = title[:57] + '...'
        docx_path = txt.with_suffix('.docx')
        n_paras = generate_word(txt, docx_path, title, creator_name)
        print(f'  ✓ {title[:50]}: {n_paras} 段')
        docs.append({**task, 'docx_path': docx_path, 'title': title})
    
    # Step 5: 上传 IMA
    print(f'\n[5/5] 上传 IMA ({len(docs)} 条)...')
    uploaded = 0
    for doc in docs:
        docx = doc['docx_path']
        fname = docx.name
        print(f'  上传: {fname[:60]}... ', end='', flush=True)
        ok, result = upload_to_ima(str(docx), fname)
        if ok:
            print(f'✅ {result}')
            uploaded += 1
        else:
            print(f'❌ {result}')
    
    total_time = time.time() - t0
    result = {
        'batch_num': batch_num,
        'downloaded': len(videos),
        'skipped': len(skipped),
        'transcribed': len(transcribed),
        'uploaded': uploaded,
        'total_time': total_time,
        'timestamp': datetime.now().isoformat(),
        'skipped_files': [s['video'].name for s in skipped],
        'transcribed_files': [t['video'].name for t in transcribed]
    }
    
    # 保存进度
    progress = load_progress()
    progress.setdefault('batches', []).append(result)
    progress['creator'] = creator_name
    progress['sec_uid'] = sec_uid
    progress['last_updated'] = datetime.now().isoformat()
    progress['total_uploaded'] = progress.get('total_uploaded', 0) + uploaded
    save_progress(progress)
    
    print(f'\n  批次 {batch_num} 完成!')
    print(f'  下载: {len(videos)} | 跳过: {len(skipped)} | 转写: {len(transcribed)} | 入库: {uploaded}')
    print(f'  总耗时: {total_time:.0f}s ({total_time/60:.1f}min)')
    
    return result


def main():
    parser = argparse.ArgumentParser(description='批量抖音博主文案提取')
    parser.add_argument('--creator', required=True, help='博主名称')
    parser.add_argument('--sec-uid', required=True, help='博主 sec_uid')
    parser.add_argument('--batch-size', type=int, default=10, help='每批数量')
    parser.add_argument('--start-batch', type=int, default=1, help='从第几批开始')
    parser.add_argument('--output-dir', default=str(WORK_DIR / 'douyin_output' / 'batch'), help='输出根目录')
    
    args = parser.parse_args()
    
    output_base = Path(args.output_dir)
    batch_num = args.start_batch
    
    result = process_batch(
        args.sec_uid, args.creator, args.batch_size, batch_num, output_base
    )
    
    # 打印汇总
    print(f'\n{"="*60}')
    print(f'  批次 {batch_num} 汇总')
    print(f'{"="*60}')
    print(f'  下载视频: {result["downloaded"]}')
    print(f'  跳过(图文): {result["skipped"]}')
    print(f'  转写成功: {result["transcribed"]}')
    print(f'  入库成功: {result["uploaded"]}')
    print(f'  总耗时: {result["total_time"]:.0f}s')
    print(f'  进度文件: {PROGRESS_FILE}')


if __name__ == '__main__':
    main()
